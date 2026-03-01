# -*- coding: utf-8 -*-
"""
数据导出模块
支持导出Excel、Word、PDF格式的报告和数据
"""
import os
import pandas as pd
from datetime import datetime
from typing import Optional, Dict, Any
import json

class DataExporter:
    """数据导出器"""
    
    def __init__(self, output_dir: str = "data/exports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export_to_excel(self, df: pd.DataFrame, filename: str = None, 
                        include_charts: bool = True) -> str:
        """
        导出数据到Excel文件
        
        Args:
            df: 数据框
            filename: 文件名（不含扩展名）
            include_charts: 是否包含图表工作表
        
        Returns:
            导出文件路径
        """
        if filename is None:
            filename = f"舆情数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        output_path = os.path.join(self.output_dir, f"{filename}.xlsx")
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='原始数据', index=False)
            
            if include_charts and len(df) > 0:
                summary_data = self._generate_summary(df)
                summary_df = pd.DataFrame([summary_data])
                summary_df.to_excel(writer, sheet_name='数据摘要', index=False)
                
                if 'platform' in df.columns:
                    platform_stats = df.groupby('platform').agg({
                        'csi_score': ['mean', 'count'],
                        'polarity_label': lambda x: (x == '积极').mean() * 100
                    }).round(2)
                    platform_stats.columns = ['平均CSI', '评论数', '积极率(%)']
                    platform_stats.to_excel(writer, sheet_name='平台统计')
                
                if 'polarity_label' in df.columns:
                    sentiment_stats = df['polarity_label'].value_counts()
                    sentiment_df = pd.DataFrame({
                        '情感类型': sentiment_stats.index,
                        '数量': sentiment_stats.values,
                        '占比(%)': (sentiment_stats.values / len(df) * 100).round(2)
                    })
                    sentiment_df.to_excel(writer, sheet_name='情感分布', index=False)
        
        return output_path
    
    def export_to_csv(self, df: pd.DataFrame, filename: str = None) -> str:
        """导出数据到CSV文件"""
        if filename is None:
            filename = f"舆情数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        output_path = os.path.join(self.output_dir, f"{filename}.csv")
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        
        return output_path
    
    def export_to_json(self, df: pd.DataFrame, filename: str = None) -> str:
        """导出数据到JSON文件"""
        if filename is None:
            filename = f"舆情数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        output_path = os.path.join(self.output_dir, f"{filename}.json")
        
        export_data = {
            "export_time": datetime.now().isoformat(),
            "total_records": len(df),
            "summary": self._generate_summary(df),
            "data": df.to_dict(orient='records')[:1000]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def export_report_to_word(self, df: pd.DataFrame, report_content: str, 
                               filename: str = None) -> str:
        """
        导出报告到Word文档
        
        Args:
            df: 数据框
            report_content: 报告内容（Markdown格式）
            filename: 文件名
        
        Returns:
            导出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("⚠️ 需要安装python-docx库: pip install python-docx")
            return None
        
        if filename is None:
            filename = f"舆情报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        output_path = os.path.join(self.output_dir, f"{filename}.docx")
        
        doc = Document()
        
        title = doc.add_heading('上海迪士尼舆情分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        doc.add_paragraph(f"数据样本: {len(df)} 条评论")
        
        doc.add_heading('数据摘要', level=1)
        summary = self._generate_summary(df)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        for key, value in summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        doc.add_heading('详细报告', level=1)
        
        lines = report_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        
        return output_path
    
    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成数据摘要"""
        summary = {
            "总评论数": len(df),
            "数据平台数": df['platform'].nunique() if 'platform' in df.columns else 0,
        }
        
        if 'csi_score' in df.columns:
            summary["平均CSI得分"] = f"{df['csi_score'].mean():.1f}"
            summary["最高CSI得分"] = f"{df['csi_score'].max():.1f}"
            summary["最低CSI得分"] = f"{df['csi_score'].min():.1f}"
        
        if 'polarity_label' in df.columns:
            sentiment_counts = df['polarity_label'].value_counts()
            for label in ['积极', '中性', '消极']:
                if label in sentiment_counts:
                    pct = sentiment_counts[label] / len(df) * 100
                    summary[f"{label}评价占比"] = f"{pct:.1f}%"
        
        if 'urgency_score' in df.columns:
            high_urgency = len(df[df['urgency_score'] >= 7])
            summary["高紧急度评论数"] = high_urgency
        
        return summary

class AlertManager:
    """舆情预警管理器"""
    
    def __init__(self, output_dir: str = "data/alerts"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.alerts = []
        self.alert_history_file = os.path.join(output_dir, "alert_history.json")
        self._load_history()
    
    def _load_history(self):
        """加载历史预警"""
        if os.path.exists(self.alert_history_file):
            try:
                with open(self.alert_history_file, 'r', encoding='utf-8') as f:
                    self.alerts = json.load(f)
            except:
                self.alerts = []
    
    def _save_history(self):
        """保存预警历史"""
        with open(self.alert_history_file, 'w', encoding='utf-8') as f:
            json.dump(self.alerts[-100:], f, ensure_ascii=False, indent=2)
    
    def check_alerts(self, df: pd.DataFrame) -> list:
        """
        检查数据中的预警情况
        
        Args:
            df: 分析后的数据框
        
        Returns:
            预警列表
        """
        new_alerts = []
        
        if len(df) == 0:
            return new_alerts
        
        if 'polarity_label' in df.columns:
            neg_rate = (df['polarity_label'] == '消极').mean()
            if neg_rate > 0.4:
                alert = {
                    "type": "high_negative_rate",
                    "level": "danger" if neg_rate > 0.5 else "warning",
                    "message": f"消极评价比例过高: {neg_rate*100:.1f}%",
                    "value": neg_rate,
                    "timestamp": datetime.now().isoformat()
                }
                new_alerts.append(alert)
        
        if 'csi_score' in df.columns:
            avg_csi = df['csi_score'].mean()
            if avg_csi < 50:
                alert = {
                    "type": "low_csi",
                    "level": "danger" if avg_csi < 40 else "warning",
                    "message": f"CSI满意度指数过低: {avg_csi:.1f}",
                    "value": avg_csi,
                    "timestamp": datetime.now().isoformat()
                }
                new_alerts.append(alert)
        
        if 'urgency_score' in df.columns:
            high_urgency = df[df['urgency_score'] >= 8]
            if len(high_urgency) > 0:
                alert = {
                    "type": "high_urgency",
                    "level": "danger",
                    "message": f"发现 {len(high_urgency)} 条高紧急度评论需要立即处理",
                    "value": len(high_urgency),
                    "timestamp": datetime.now().isoformat()
                }
                new_alerts.append(alert)
        
        for alert in new_alerts:
            self.alerts.append(alert)
        
        if new_alerts:
            self._save_history()
        
        return new_alerts
    
    def get_recent_alerts(self, limit: int = 10) -> list:
        """获取最近的预警"""
        return self.alerts[-limit:] if self.alerts else []
    
    def get_alert_summary(self) -> Dict[str, int]:
        """获取预警统计"""
        summary = {
            "total": len(self.alerts),
            "danger": len([a for a in self.alerts if a.get('level') == 'danger']),
            "warning": len([a for a in self.alerts if a.get('level') == 'warning'])
        }
        return summary

class DataComparator:
    """数据对比分析器"""
    
    def __init__(self):
        self.comparison_results = {}
    
    def compare_periods(self, df1: pd.DataFrame, df2: pd.DataFrame, 
                        label1: str = "时期1", label2: str = "时期2") -> Dict[str, Any]:
        """
        对比两个时期的数据
        
        Args:
            df1: 时期1的数据
            df2: 时期2的数据
            label1: 时期1标签
            label2: 时期2标签
        
        Returns:
            对比结果
        """
        results = {
            "period1": label1,
            "period2": label2,
            "metrics": {}
        }
        
        if 'csi_score' in df1.columns and 'csi_score' in df2.columns:
            csi1 = df1['csi_score'].mean()
            csi2 = df2['csi_score'].mean()
            results["metrics"]["CSI满意度"] = {
                label1: round(csi1, 1),
                label2: round(csi2, 1),
                "变化": round(csi2 - csi1, 1),
                "变化率": f"{(csi2-csi1)/csi1*100:.1f}%" if csi1 != 0 else "N/A"
            }
        
        if 'polarity_label' in df1.columns and 'polarity_label' in df2.columns:
            for label in ['积极', '中性', '消极']:
                rate1 = (df1['polarity_label'] == label).mean() * 100
                rate2 = (df2['polarity_label'] == label).mean() * 100
                results["metrics"][f"{label}评价占比"] = {
                    label1: round(rate1, 1),
                    label2: round(rate2, 1),
                    "变化": round(rate2 - rate1, 1)
                }
        
        results["metrics"]["评论数量"] = {
            label1: len(df1),
            label2: len(df2),
            "变化": len(df2) - len(df1)
        }
        
        self.comparison_results = results
        return results
    
    def compare_platforms(self, df: pd.DataFrame) -> Dict[str, Any]:
        """
        对比不同平台的数据
        
        Args:
            df: 包含platform列的数据框
        
        Returns:
            平台对比结果
        """
        if 'platform' not in df.columns:
            return {"error": "数据中缺少platform列"}
        
        results = {}
        
        for platform in df['platform'].unique():
            platform_df = df[df['platform'] == platform]
            stats = {
                "评论数": len(platform_df)
            }
            
            if 'csi_score' in platform_df.columns:
                stats["平均CSI"] = round(platform_df['csi_score'].mean(), 1)
            
            if 'polarity_label' in platform_df.columns:
                stats["积极率"] = round(
                    (platform_df['polarity_label'] == '积极').mean() * 100, 1
                )
            
            results[platform] = stats
        
        return results

if __name__ == "__main__":
    print("数据导出模块加载完成")
